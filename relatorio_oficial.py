import io
import os
from datetime import datetime


COORDENADORA = "Rosiane Rodrigues Shiozaki"
SIBEC_CATEGORIAS = [
    "Consulta de Benefício",
    "Bloqueio de Benefício",
    "Desbloqueio de Benefício",
    "Cancelamento de Benefício",
    "Reversão de Cancelamento",
    "Revisão Cadastral",
    "Averiguação Cadastral",
]
TIPOS_CADUNICO = [
    "CadÚnico para BPC - 1ª Vez (Idoso)",
    "CadÚnico para BPC - 1ª Vez (PCD)",
    "CadÚnico para BPC - Atualização (Idoso)",
    "CadÚnico para BPC - Atualização (PCD)",
    "Carteira do Idoso Emitida",
    "Comprovante de Cadastro",
    "Consulta Cadastro Único",
    "Encaminhamentos",
    "Escuta Qualificada",
    "Benefício Eventual",
    "CadÚnico Unipessoal",
    "Exclusão de membros",
    "Inclusão de membros",
    "Inscrição do CadÚnico",
    "Recadastramento",
    "Tarifa Social de Energia Elétrica",
    "Transferência de Município",
    "Troca de RF",
    "Visita Domiciliar",
]

# Mapeamento: coluna abreviada → tipo(s) do sistema
COLUNAS_RMA = [
    ("RECAD",          ["Recadastramento"]),
    ("BPC 1ª IDOSO",   ["CadÚnico para BPC - 1ª Vez (Idoso)"]),
    ("BPC 1ª PCD",     ["CadÚnico para BPC - 1ª Vez (PCD)", "CadÚnico para BPC - 1ª Vez"]),
    ("BPC ATU IDOSO",  ["CadÚnico para BPC - Atualização (Idoso)"]),
    ("BPC ATU PCD",    ["CadÚnico para BPC - Atualização (PCD)", "CadÚnico para BPC - Atualização"]),
    ("INC",            ["Inclusão de membros"]),
    ("EXC",            ["Exclusão de membros"]),
    ("ENCAMINH.",      ["Encaminhamentos", "Atendimento Técnico / Parecer Social"]),
    ("ESCUTA QUALIF.", ["Escuta Qualificada"]),
    ("BENEF. EVENT.",  ["Benefício Eventual"]),
    ("TSEE",           ["Tarifa Social de Energia Elétrica"]),
    ("TRANSF.\nMUN",   ["Transferência de Município"]),
    ("CART.\nIDOSO",   ["Carteira do Idoso Emitida"]),
    ("UNIPESSOAL",     ["CadÚnico Unipessoal"]),
    ("FOLHA\nRESUMO",  ["Folha de Pagamento (SIBEC)"]),
    ("INFO",           ["Comprovante de Cadastro", "Consulta Cadastro Único", "Consulta SIBEC"]),
    ("Troca\nde RF",   ["Troca de RF"]),
    ("vis.\ndom",      ["Visita Domiciliar"]),
]

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _fmt_data(data):
    return datetime.strptime(data, "%Y-%m-%d").strftime("%d/%m/%Y")


def _total_tipo(quant, tipo):
    return sum(quant.get(tipo, {}).values())


def _valor_quant(quant, tipo, entrevistador):
    return quant.get(tipo, {}).get(entrevistador, 0)


def _mapear_sibec(atendimentos):
    totais = {categoria: 0 for categoria in SIBEC_CATEGORIAS}
    for atendimento in atendimentos:
        for tipo in atendimento["tipos"].split("|"):
            if tipo in ("Consulta SIBEC", "Folha de Pagamento (SIBEC)"):
                totais["Consulta de Benefício"] += 1
            elif tipo == "Bloqueio de Benefício":
                totais["Bloqueio de Benefício"] += 1
            elif tipo == "Desbloqueio de Benefício":
                totais["Desbloqueio de Benefício"] += 1
            elif tipo == "Reversão de Cancelamento":
                totais["Reversão de Cancelamento"] += 1
    return totais


def _total_sibec(totais):
    return sum(totais.values())


def _asset(assets_dir, filename):
    path = os.path.join(assets_dir, filename)
    return path if os.path.exists(path) else None


import re

def _clean_html_for_pdf(html_text):
    if not html_text:
        return ""
    s = html_text
    s = re.sub(r'<p[^>]*>', '', s)
    s = re.sub(r'</p>', '<br/><br/>', s)
    s = re.sub(r'<strong>', '<b>', s)
    s = re.sub(r'</strong>', '</b>', s)
    s = re.sub(r'<em>', '<i>', s)
    s = re.sub(r'</em>', '</i>', s)
    s = re.sub(r'<li[^>]*>', '• ', s)
    s = re.sub(r'</li>', '<br/>', s)
    s = re.sub(r'</?(ul|ol|h[1-6]|blockquote)[^>]*>', '', s)
    s = re.sub(r'(<br\s*/?>\s*){3,}', '<br/><br/>', s)
    return s.strip()


def _add_html_to_pdf_story(story, html_text, style):
    from reportlab.platypus import Paragraph
    clean = _clean_html_for_pdf(html_text)
    if clean:
        for block in clean.split('<br/><br/>'):
            txt = block.strip()
            if txt:
                story.append(Paragraph(txt, style))


def _add_html_to_docx(doc, html_text):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    if not html_text:
        return

    if '<' not in html_text or '>' not in html_text:
        for p_str in html_text.split('\n\n'):
            if p_str.strip():
                p = doc.add_paragraph(p_str.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return

    blocks = re.findall(r'<(p|li|h[1-6]|blockquote)[^>]*>(.*?)</\1>', html_text, re.DOTALL | re.IGNORECASE)
    if not blocks:
        blocks = [('p', b) for b in re.split(r'<br\s*/?>', html_text) if b.strip()]

    for tag, content in blocks:
        is_list = (tag.lower() == 'li')
        p = doc.add_paragraph(style='List Bullet' if is_list else 'Normal')
        if not is_list:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        tokens = re.split(r'(</?(?:strong|b|em|i|u|br)[^>]*>)', content, flags=re.IGNORECASE)
        is_b = False
        is_i = False
        is_u = False

        for tok in tokens:
            ltok = tok.lower()
            if ltok in ('<strong>', '<b>'):
                is_b = True
            elif ltok in ('</strong>', '</b>'):
                is_b = False
            elif ltok in ('<em>', '<i>'):
                is_i = True
            elif ltok in ('</em>', '</i>'):
                is_i = False
            elif ltok in ('<u>',):
                is_u = True
            elif ltok in ('</u>',):
                is_u = False
            elif ltok.startswith('<br'):
                p.add_run('\n')
            elif not tok.startswith('<'):
                text = re.sub(r'&[a-z]+;', ' ', tok)
                if text:
                    run = p.add_run(text)
                    run.bold = is_b
                    run.italic = is_i
                    run.underline = is_u
                    run.font.size = Pt(9)


def _dependency_error(package):
    raise RuntimeError(
        f"Dependência ausente: {package}. Execute: pip install -r requirements.txt"
    )


def _add_heading(doc, text):
    from docx.shared import Pt, RGBColor

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 84, 42)


def _style_header_cell_docx(cell):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "00542A")
    tc_pr.append(shd)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(255, 255, 255)


def _style_sibec_header_cell_docx(cell):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "783DB2")
    tc_pr.append(shd)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(255, 255, 255)


def _style_body_cell_docx(cell, bold=False):
    from docx.shared import Pt

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = bold
            run.font.size = Pt(7.5)


def _add_assinaturas_docx(doc, assinatura_nome, coordenadora=None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    coord_nome = coordenadora or COORDENADORA
    table = doc.add_table(rows=3, cols=2)
    table.autofit = True
    linhas = [
        ("___________________________________", "___________________________________"),
        ("COORD. DO CADASTRO ÚNICO/PBF", "ENTREVISTADOR"),
        (coord_nome.upper(), assinatura_nome.upper()),
    ]
    for row_idx, valores in enumerate(linhas):
        for col_idx, valor in enumerate(valores):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = valor
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _p(text, style):
    from reportlab.platypus import Paragraph
    from xml.sax.saxutils import escape

    return Paragraph(escape(str(text)), style)


def _table_style():
    from reportlab.lib import colors
    from reportlab.platypus import TableStyle

    return TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#00542A")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (1, 1), (-1, -1), "CENTER"),
        ("BACKGROUND", (-1, 1), (-1, -1), colors.HexColor("#F0FDF4")),
    ])


def _sibec_table_style():
    from reportlab.lib import colors
    from reportlab.platypus import TableStyle

    return TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#783DB2")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#E9D5FF")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (1, 1), (-1, -1), "CENTER"),
        ("BACKGROUND", (-1, 1), (-1, -1), colors.HexColor("#F3E8FF")),
    ])


def _assinaturas_pdf(assinatura_nome, style, coordenadora=None):
    from reportlab.lib.units import cm
    from reportlab.platypus import Table, TableStyle

    coord_nome = coordenadora or COORDENADORA
    data = [
        ["___________________________________", "___________________________________"],
        ["COORD. DO CADASTRO ÚNICO/PBF", "ENTREVISTADOR"],
        [coord_nome.upper(), assinatura_nome.upper()],
    ]
    table = Table(data, colWidths=[13.5 * cm, 13.5 * cm])
    table.setStyle(TableStyle([
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
    ]))
    return table


# ---------------------------------------------------------------------------
# PDF helpers (quantitative table, SIBEC table, details table)
# ---------------------------------------------------------------------------

def _header_style_pdf(style):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.styles import ParagraphStyle

    return ParagraphStyle(
        "HeaderStylePDF",
        parent=style,
        fontName="Helvetica-Bold",
        fontSize=8,
        textColor=colors.white,
        alignment=TA_CENTER,
    )


def _rma_table_pdf(quant, style):
    """Tabela horizontal no formato RMA — abreviações no cabeçalho, totais na linha abaixo."""
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import Table, TableStyle

    hstyle = _header_style_pdf(style)
    headers = [_p(col, hstyle) for col, _ in COLUNAS_RMA]
    valores = []
    for col, tipos in COLUNAS_RMA:
        total = sum(sum(quant.get(t, {}).values()) for t in tipos)
        valores.append(_p(f"{total:02d}", style))

    data = [headers, valores]
    col_w = 27.8 * cm / len(COLUNAS_RMA)
    table = Table(data, colWidths=[col_w] * len(COLUNAS_RMA))
    table.setStyle(_table_style())
    return table


def _quant_table_pdf(quant, style):
    """Tabela de 2 colunas para PDF: Tipo de Atendimento | Quantidade."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import Table, TableStyle

    hstyle = _header_style_pdf(style)
    total_style = ParagraphStyle("TotalPDF", parent=style, fontName="Helvetica-Bold", fontSize=8.5, textColor=colors.white, alignment=TA_CENTER)

    data = [[_p("Tipo de Atendimento", hstyle), _p("Quantidade", hstyle)]]
    total_geral = 0
    for tipo in TIPOS_CADUNICO:
        qtd = sum(quant.get(tipo, {}).values())
        total_geral += qtd
        data.append([_p(tipo, style), _p(str(qtd) if qtd > 0 else "-", style)])

    data.append([_p("TOTAL", total_style), _p(str(total_geral), total_style)])

    first_width = 21.0 * cm
    second_width = 6.8 * cm
    table = Table(data, repeatRows=1, colWidths=[first_width, second_width])
    style_cmds = _table_style().getCommands()
    style_cmds.extend([
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#00542A")),
        ("TEXTCOLOR", (0, -1), (-1, -1), colors.white),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
    ])
    table.setStyle(TableStyle(style_cmds))
    return table


def _sibec_table_pdf(sibec, style):
    """Builds the SIBEC quantitative table (5 visible columns)."""
    from reportlab.lib.units import cm
    from reportlab.platypus import Table

    hstyle = _header_style_pdf(style)
    visible = [
        "Consulta de Benefício",
        "Bloqueio de Benefício",
        "Desbloqueio de Benefício",
        "Reversão de Cancelamento",
        "Cancelamento de Benefício",
    ]
    extra = ["Revisão Cadastral", "Averiguação Cadastral"]
    headers = [_p(c, hstyle) for c in visible] + [_p(c, hstyle) for c in extra]
    values = [f"{sibec.get(c, 0):02d}" for c in visible] + ["00", "00"]
    data = [headers, values]
    col_w = 27.8 * cm / len(headers)
    table = Table(data, repeatRows=1, colWidths=[col_w] * len(headers))
    table.setStyle(_sibec_table_style())
    return table


def _detalhes_table_pdf(atendimentos, style):
    """Builds the detailed attendance table."""
    from reportlab.lib.units import cm
    from reportlab.platypus import Table

    hstyle = _header_style_pdf(style)
    data = [[_p(h, hstyle) for h in ["Data", "Entrevistador", "CPF", "Nome do RF", "Origem", "Tipos de Atendimento"]]]
    if not atendimentos:
        data.append([_p("Nenhum atendimento no período selecionado.", style), "", "", "", "", ""])
    else:
        for at in atendimentos:
            data.append([
                _fmt_data(at["data"]),
                _p(at["entrevistador"], style),
                at["cpf"],
                _p(at["nome_rf"], style),
                _p(at["origem"], style),
                _p(at["tipos"].replace("|", ", "), style),
            ])
    table = Table(data, repeatRows=1, colWidths=[1.8 * cm, 4.3 * cm, 3.0 * cm, 5.0 * cm, 3.2 * cm, 10.5 * cm])
    table.setStyle(_table_style())
    return table


# ---------------------------------------------------------------------------
# 1. criar_pdf_oficial — Relatório Quantitativo Oficial (PDF, sem lista detalhada)
# ---------------------------------------------------------------------------

def criar_pdf_oficial(
    *,
    mes,
    mes_nome,
    atendimentos,
    quant,
    entrevistadores,
    tipos_atendimento,
    total_geral,
    assinatura_nome,
    assets_dir,
    config=None,
):
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Image as RLImage
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError:
        _dependency_error("reportlab")

    output = io.BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        rightMargin=0.9 * cm,
        leftMargin=0.9 * cm,
        topMargin=0.8 * cm,
        bottomMargin=0.8 * cm,
    )
    styles = getSampleStyleSheet()
    normal = ParagraphStyle("NormalCad", parent=styles["Normal"], fontName="Helvetica", fontSize=8.2, leading=10)
    centered = ParagraphStyle("CenteredCad", parent=normal, alignment=TA_CENTER)
    justified = ParagraphStyle("JustifiedCad", parent=normal, alignment=TA_JUSTIFY)
    heading = ParagraphStyle(
        "HeadingCad",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=10,
        textColor=colors.HexColor("#1F4E79"),
        spaceBefore=8,
        spaceAfter=5,
    )
    title = ParagraphStyle(
        "TitleCad",
        parent=styles["Title"],
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
        fontSize=13,
        textColor=colors.HexColor("#1F4E79"),
        leading=15,
    )

    story = []

    # 1. Header image
    header = _asset(assets_dir, "image3.png")
    if header:
        img = RLImage(header, width=25.5 * cm, height=3.4 * cm)
        img._restrictSize(25.5 * cm, 3.4 * cm)
        story.append(img)
        story.append(Spacer(1, 4))

    # 2. Title
    story.append(Paragraph(f"RELATÓRIO DO CADASTRO ÚNICO/PBF {mes_nome.upper()}", title))

    # 3. Subtitle
    story.append(Paragraph("Departamento do Cadastro Único e Programa Bolsa Família", centered))
    story.append(Spacer(1, 6))

    # 4. Logos avulsas removidas conforme solicitado (o cabeçalho já possui todos os logotipos embutidos)

    # 5. Section 1 - Identificação
    cfg = config or {}
    coordenadora = cfg.get('coordenadora', COORDENADORA)
    story.append(Paragraph("1. - Identificação", heading))
    for linha in [
        f"Nome: {cfg.get('setor_nome', 'Setor do Cadastro Único /PBF')}",
        f"Endereço: {cfg.get('endereco', 'Bairro: Campina - Rua Bruno de Menezes s/n')}",
        f"E-mail: {cfg.get('email_setor', 'setascadastrounico@gmail.com')}",
        f"Território de Abrangência: {cfg.get('territorio', 'Sede do Município e Distrito de Quatro Bocas')}",
        f"Coordenadora: {coordenadora}",
    ]:
        story.append(Paragraph(linha, normal))

    if cfg.get('texto_identificacao'):
        _add_html_to_pdf_story(story, cfg['texto_identificacao'], justified)

    # 6. Section 2 - Apresentação
    story.append(Paragraph("2. - Apresentação", heading))
    if cfg.get('texto_apresentacao'):
        _add_html_to_pdf_story(story, cfg['texto_apresentacao'], justified)
    else:
        story.append(Paragraph(
            "A equipe do Cadastro Único realiza atendimentos, orientações e encaminhamentos "
            "relacionados aos serviços do Cadastro Único e Programa Bolsa Família. ",
            justified,
        ))
    story.append(Paragraph(
        f"Abaixo contém a quantificação dos atendimentos realizados por este setor no mês "
        f"de {mes_nome}, totalizando {total_geral:02d} atendimento"
        f"{'s' if total_geral != 1 else ''}.",
        justified,
    ))
    story.append(Spacer(1, 6))

    # 7. Table 1 — CadÚnico
    story.append(Paragraph("Tabela 1 – Quantitativo Cadúnico", heading))
    story.append(_quant_table_pdf(quant, normal))
    story.append(Spacer(1, 8))

    # 8. Table 2 — SIBEC quantitative
    sibec = _mapear_sibec(atendimentos)
    story.append(Paragraph("Tabela 2 — Quantitativo SIBEC", heading))
    story.append(_sibec_table_pdf(sibec, normal))
    story.append(Spacer(1, 24))

    # 9. Signature block
    story.append(_assinaturas_pdf(assinatura_nome, normal, coordenadora))

    # 10. Footer image (imagem de faixa colorida enviada)
    footer = _asset(assets_dir, "image4.png")
    if footer:
        story.append(Spacer(1, 0.5 * cm))
        img = RLImage(footer, width=25.5 * cm, height=0.6 * cm)
        img._restrictSize(25.5 * cm, 0.6 * cm)
        story.append(img)

    # 11. Texto do Rodapé (0,5 cm abaixo da imagem do rodapé)
    rodape_txt = cfg.get('rodape', '')
    if rodape_txt:
        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph(rodape_txt.replace('\n', '<br/>'), centered))

    doc.build(story)
    output.seek(0)
    return output


# ---------------------------------------------------------------------------
# 2. criar_pdf_registro_detalhado — Registro Detalhado de Atendimentos (PDF)
# ---------------------------------------------------------------------------

def criar_pdf_registro_detalhado(
    *,
    mes,
    mes_nome,
    atendimentos,
    assinatura_nome,
    assets_dir,
    quant=None,
    total_geral=0,
):
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Image as RLImage
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError:
        _dependency_error("reportlab")

    # Ordenar cronologicamente do início do mês até o final do mês (dia 1 ao 31)
    atendimentos_ord = sorted(atendimentos, key=lambda x: (x["data"], dict(x).get("id", 0)))

    output = io.BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        rightMargin=0.9 * cm,
        leftMargin=0.9 * cm,
        topMargin=0.8 * cm,
        bottomMargin=0.8 * cm,
    )
    styles = getSampleStyleSheet()
    normal = ParagraphStyle("NormalReg", parent=styles["Normal"], fontName="Helvetica", fontSize=8.2, leading=10)
    centered = ParagraphStyle("CenteredReg", parent=normal, alignment=TA_CENTER)
    title = ParagraphStyle(
        "TitleReg",
        parent=styles["Title"],
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
        fontSize=13,
        textColor=colors.HexColor("#00542A"),
        leading=15,
    )

    story = []

    # 1. Header image
    header = _asset(assets_dir, "image3.png")
    if header:
        img = RLImage(header, width=25.5 * cm, height=3.4 * cm)
        img._restrictSize(25.5 * cm, 3.4 * cm)
        story.append(img)
        story.append(Spacer(1, 4))

    # 2. Title
    story.append(Paragraph(f"REGISTRO DETALHADO DE ATENDIMENTOS — {mes_nome.upper()}", title))
    story.append(Spacer(1, 8))

    # 3. Detailed attendance table (chronological order)
    story.append(_detalhes_table_pdf(atendimentos_ord, normal))
    story.append(Spacer(1, 16))

    # 4. Resumo Quantitativo no fim do documento
    if quant:
        story.append(Paragraph(f"RESUMO QUANTITATIVO DE ATENDIMENTOS — {mes_nome.upper()}", title))
        story.append(Spacer(1, 8))
        story.append(_quant_table_pdf(quant, normal))
        story.append(Spacer(1, 16))

    # 5. Signature block
    story.append(_assinaturas_pdf(assinatura_nome, normal))

    # 6. Footer image
    footer = _asset(assets_dir, "image4.png")
    if footer:
        story.append(Spacer(1, 12))
        img = RLImage(footer, width=25.5 * cm, height=1.5 * cm)
        img._restrictSize(25.5 * cm, 1.5 * cm)
        story.append(img)

    doc.build(story)
    output.seek(0)
    return output


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _add_rma_table_docx(doc, quant):
    """Tabela horizontal RMA para Word — abreviações no cabeçalho, totais na linha abaixo."""
    table = doc.add_table(rows=2, cols=len(COLUNAS_RMA))
    table.style = "Table Grid"
    for idx, (col, tipos) in enumerate(COLUNAS_RMA):
        # Cabeçalho
        cell = table.rows[0].cells[idx]
        cell.text = col
        _style_header_cell_docx(cell)
        # Valor
        total = sum(sum(quant.get(t, {}).values()) for t in tipos)
        table.rows[1].cells[idx].text = f"{total:02d}"
        _style_body_cell_docx(table.rows[1].cells[idx], bold=True)


def _add_quant_table_docx(doc, quant):
    """Tabela de 2 colunas para Word: Tipo de Atendimento | Quantidade."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = ["Tipo de Atendimento", "Quantidade"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        _style_header_cell_docx(cell)

    total_geral = 0
    for tipo in TIPOS_CADUNICO:
        row = table.add_row().cells
        row[0].text = tipo
        _style_body_cell_docx(row[0], bold=False)
        qtd = sum(quant.get(tipo, {}).values())
        total_geral += qtd
        row[1].text = str(qtd) if qtd > 0 else "-"
        _style_body_cell_docx(row[1], bold=bool(qtd))

    total = table.add_row().cells
    total[0].text = "TOTAL"
    _style_header_cell_docx(total[0])
    total[1].text = str(total_geral)
    _style_header_cell_docx(total[1])


def _add_sibec_table_docx(doc, sibec):
    """SIBEC table with 5 main cols + 2 showing 00."""
    visible = [
        "Consulta de Benefício",
        "Bloqueio de Benefício",
        "Desbloqueio de Benefício",
        "Reversão de Cancelamento",
        "Cancelamento de Benefício",
        "Revisão Cadastral",
        "Averiguação Cadastral",
    ]
    table = doc.add_table(rows=2, cols=len(visible))
    table.style = "Table Grid"
    for idx, categoria in enumerate(visible):
        table.rows[0].cells[idx].text = categoria
        _style_sibec_header_cell_docx(table.rows[0].cells[idx])
        if categoria in ("Revisão Cadastral", "Averiguação Cadastral"):
            table.rows[1].cells[idx].text = "00"
        else:
            table.rows[1].cells[idx].text = f"{sibec.get(categoria, 0):02d}"
        _style_body_cell_docx(table.rows[1].cells[idx], bold=True)


def _add_detalhes_table_docx(doc, atendimentos):
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Data", "Entrevistador", "CPF", "Nome do RF", "Origem", "Tipos de Atendimento"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
        _style_header_cell_docx(table.rows[0].cells[idx])

    if not atendimentos:
        row = table.add_row().cells
        row[0].text = "Nenhum atendimento no período selecionado."
        _style_body_cell_docx(row[0])
        return

    for atendimento in atendimentos:
        row = table.add_row().cells
        valores = [
            _fmt_data(atendimento["data"]),
            atendimento["entrevistador"],
            atendimento["cpf"],
            atendimento["nome_rf"],
            atendimento["origem"],
            atendimento["tipos"].replace("|", ", "),
        ]
        for idx, valor in enumerate(valores):
            row[idx].text = str(valor)
            _style_body_cell_docx(row[idx])


# ---------------------------------------------------------------------------
# 3. criar_docx_oficial — Word version of the quantitative report
# ---------------------------------------------------------------------------

def criar_docx_oficial(
    *,
    mes,
    mes_nome,
    atendimentos,
    quant,
    entrevistadores,
    tipos_atendimento,
    total_geral,
    assinatura_nome,
    assets_dir,
    config=None,
):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        _dependency_error("python-docx")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    header = _asset(assets_dir, "image3.png")
    cadunico = _asset(assets_dir, "image1.png")
    bolsa = _asset(assets_dir, "image2.png")
    footer = _asset(assets_dir, "image4.png")

    # 1. Header image
    if header:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(header, width=Inches(6.6))

    # 2. Title
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run(f"RELATÓRIO DO CADASTRO ÚNICO/PBF {mes_nome.upper()}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(31, 78, 121)

    # 3. Subtitle
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run("Departamento do Cadastro Único e Programa Bolsa Família")
    run.bold = True
    run.font.size = Pt(10)

    # 4. Section 1 - Identificação
    cfg = config or {}
    coordenadora = cfg.get('coordenadora', COORDENADORA)
    _add_heading(doc, "- Identificação")
    for linha in [
        f"Nome: {cfg.get('setor_nome', 'Setor do Cadastro Único /PBF')}",
        f"Endereço: {cfg.get('endereco', 'Bairro: Campina - Rua Bruno de Menezes s/n')}",
        f"E-mail: {cfg.get('email_setor', 'setascadastrounico@gmail.com')}",
        f"Território de Abrangência: {cfg.get('territorio', 'Sede do Município e Distrito de Quatro Bocas')}",
        f"Coordenadora: {coordenadora}",
    ]:
        doc.add_paragraph(linha)

    if cfg.get('texto_identificacao'):
        _add_html_to_docx(doc, cfg['texto_identificacao'])

    # 5. Section 2 - Apresentação
    _add_heading(doc, "- Apresentação")
    if cfg.get('texto_apresentacao'):
        _add_html_to_docx(doc, cfg['texto_apresentacao'])
    else:
        apresentacao = (
            "A equipe do Cadastro Único realiza atendimentos, orientações e encaminhamentos "
            "relacionados aos serviços do Cadastro Único e Programa Bolsa Família. "
        )
        p = doc.add_paragraph(apresentacao)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = doc.add_paragraph(
        f"Abaixo contém a quantificação dos atendimentos realizados por este setor no mês "
        f"de {mes_nome}, totalizando {total_geral:02d} atendimento"
        f"{'s' if total_geral != 1 else ''}."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 7. Table 1 — CadÚnico
    _add_heading(doc, "Tabela 1 – Quantitativo Cadúnico")
    _add_quant_table_docx(doc, quant)

    # 8. Table 2 — SIBEC quantitative
    sibec = _mapear_sibec(atendimentos)
    _add_heading(doc, "Tabela 2 — Quantitativo SIBEC")
    _add_sibec_table_docx(doc, sibec)

    doc.add_paragraph("")

    # 9. Signature block
    _add_assinaturas_docx(doc, assinatura_nome, coordenadora)

    # 10. Footer image (imagem de faixa colorida enviada)
    if footer:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(footer, width=Inches(6.6))

    # 11. Texto do Rodapé (0,5 cm abaixo da imagem do rodapé)
    rodape_txt = cfg.get('rodape', '')
    if rodape_txt:
        p_txt = doc.add_paragraph()
        p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_txt.paragraph_format.space_before = Pt(14)
        for idx, linha in enumerate(rodape_txt.split('\n')):
            if idx > 0:
                p_txt.add_run('\n')
            run = p_txt.add_run(linha)
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(100, 100, 100)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# ---------------------------------------------------------------------------
# 4. criar_docx_registro_detalhado — Word version of detailed register
# ---------------------------------------------------------------------------

def criar_docx_registro_detalhado(
    *,
    mes,
    mes_nome,
    atendimentos,
    assinatura_nome,
    assets_dir,
    quant=None,
    total_geral=0,
):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        _dependency_error("python-docx")

    # Ordenar cronologicamente do início do mês até o final do mês (dia 1 ao 31)
    atendimentos_ord = sorted(atendimentos, key=lambda x: (x["data"], dict(x).get("id", 0)))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    header = _asset(assets_dir, "image3.png")
    footer = _asset(assets_dir, "image4.png")

    # 1. Header image
    if header:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(header, width=Inches(6.6))

    # 2. Title
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run(f"REGISTRO DETALHADO DE ATENDIMENTOS — {mes_nome.upper()}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 84, 42)

    doc.add_paragraph("")

    # 3. Detailed attendance table (chronological order)
    _add_detalhes_table_docx(doc, atendimentos_ord)

    doc.add_paragraph("")

    # 4. Resumo Quantitativo no fim do documento
    if quant:
        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        srun = subtitulo.add_run(f"RESUMO QUANTITATIVO DE ATENDIMENTOS — {mes_nome.upper()}")
        srun.bold = True
        srun.font.size = Pt(12)
        srun.font.color.rgb = RGBColor(0, 84, 42)

        doc.add_paragraph("")
        _add_quant_table_docx(doc, quant)
        doc.add_paragraph("")

    # 5. Signature block
    _add_assinaturas_docx(doc, assinatura_nome)

    # 6. Footer image
    if footer:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(footer, width=Inches(6.6))

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
